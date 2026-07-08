"""校閲（レビュー）系ツール群。2ファイルの差分比較 + Wordの変更履歴・コメント。"""
import copy
import difflib
import itertools
from datetime import datetime, timezone

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from langchain_core.tools import tool

from ...atomic import atomic_save
from ...config import resolve_workspace_path
from ...services.textextract import EXTRACTORS

# 差分出力が長くなりすぎないよう、返す行数・文字数に上限を設ける
_MAX_DIFF_LINES = 200
_MAX_DIFF_CHARS = 4000


@tool
def doc_diff(file_a: str, file_b: str) -> str:
    """2つのファイルの内容を比較し、どこがどう変わったかを差分で返す。
    「AとBの違いを教えて」「前の版と何が変わった？」に使う(file_a=元/古い版、file_b=新しい版)。
    対応形式は .docx / .xlsx / .pptx / .csv / .txt で、両方とも同じ形式である必要がある。
    書式(色・フォント等)ではなく文章・セルの値など「中身」を比べる。
    結果は「－」で始まる行が削除、「＋」で始まる行が追加を表す。この差分をもとに、
    変更点を平易な日本語で要約して報告すること(差分記号そのものはユーザーに見せない)。"""
    path_a = resolve_workspace_path(file_a, must_exist=True)
    path_b = resolve_workspace_path(file_b, must_exist=True)
    suffix = path_a.suffix.lower()
    if path_b.suffix.lower() != suffix:
        return "エラー: 同じ形式のファイル同士を指定してください(例: .docx と .docx)"
    extractor = EXTRACTORS.get(suffix)
    if extractor is None:
        return f"エラー: {suffix} は比較に対応していません(.docx / .xlsx / .pptx / .csv / .txt)"

    try:
        lines_a = extractor(str(path_a))
        lines_b = extractor(str(path_b))
    except Exception as e:  # 破損ファイル等
        return f"エラー: ファイルを読み取れませんでした: {type(e).__name__}: {e}"

    added = removed = 0
    out: list[str] = []
    for line in difflib.unified_diff(lines_a, lines_b, lineterm="", n=2):
        if line.startswith(("---", "+++")):
            continue  # ファイルヘッダー行は不要
        if line.startswith("@@"):
            out.append("……")
        elif line.startswith("+"):
            added += 1
            out.append(f"＋ {line[1:]}")
        elif line.startswith("-"):
            removed += 1
            out.append(f"－ {line[1:]}")
        else:
            out.append(f"  {line[1:] if line.startswith(' ') else line}")

    if added == 0 and removed == 0:
        return f"「{file_a}」と「{file_b}」の内容に違いはありませんでした。"

    body = "\n".join(out)
    truncated = False
    if len(out) > _MAX_DIFF_LINES or len(body) > _MAX_DIFF_CHARS:
        body = "\n".join(out[:_MAX_DIFF_LINES])[:_MAX_DIFF_CHARS]
        truncated = True

    header = f"「{file_a}」→「{file_b}」の差分(－削除 / ＋追加、追加{added}行・削除{removed}行):"
    result = f"{header}\n{body}"
    if truncated:
        result += "\n…(差分が多いため一部のみ表示。主要な変更点を優先して要約してください)"
    return result


# --- Wordの変更履歴（トラックチェンジ）・コメント -----------------------------
# python-docxには変更履歴のAPIが無いため、w:ins/w:del を直接組み立てる。
# 生成した .docx は本物のWordで「変更履歴」として承認/却下でき、プレビュー
# (docx-preview の renderChanges/renderComments)でも赤字・吹き出しで表示される。


def _open_docx(filename: str):
    path = resolve_workspace_path(filename, must_exist=True)
    if path.suffix.lower() != ".docx":
        raise ValueError("Word文書(.docx)を指定してください")
    return Document(str(path)), path


def _now_iso() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _start_revision_id(doc) -> "itertools.count":
    """既存の変更履歴IDと衝突しない採番用カウンタを返す。"""
    ids = [0]
    for el in doc.element.iter():
        if el.tag in (qn("w:ins"), qn("w:del")):
            v = el.get(qn("w:id"))
            if v and v.lstrip("-").isdigit():
                ids.append(int(v))
    return itertools.count(max(ids) + 1)


def _set_rev_attrs(el, wid: int, author: str, date: str) -> None:
    el.set(qn("w:id"), str(wid))
    el.set(qn("w:author"), author)
    el.set(qn("w:date"), date)


def _has_revision(p) -> bool:
    return p._p.find(qn("w:ins")) is not None or p._p.find(qn("w:del")) is not None


def _tracked_replace(p, new_text: str, ids, author: str, date: str) -> None:
    """段落の既存本文を「削除(取り消し線)」扱いにし、new_text を「挿入」扱いで足す。"""
    old_runs = list(p.runs)
    # 挿入する新しいラン(書式は元の先頭ランから引き継ぐ)
    new_r = OxmlElement("w:r")
    if old_runs:
        rpr = old_runs[0]._r.find(qn("w:rPr"))
        if rpr is not None:
            new_r.append(copy.deepcopy(rpr))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = new_text
    new_r.append(t)
    # 既存ランを w:del でまとめ、w:t → w:delText に変える
    if old_runs:
        del_el = OxmlElement("w:del")
        _set_rev_attrs(del_el, next(ids), author, date)
        old_runs[0]._r.addprevious(del_el)
        for r in old_runs:
            for tt in r._r.findall(qn("w:t")):
                tt.tag = qn("w:delText")
            del_el.append(r._r)
    # 挿入ランを w:ins で包んで段落末尾に置く
    ins_el = OxmlElement("w:ins")
    _set_rev_attrs(ins_el, next(ids), author, date)
    ins_el.append(new_r)
    p._p.append(ins_el)


def _tracked_delete(p, ids, author: str, date: str) -> None:
    """段落本文を「削除(取り消し線)」扱いにする。承認すると本文が消える。"""
    old_runs = list(p.runs)
    if not old_runs:
        return
    del_el = OxmlElement("w:del")
    _set_rev_attrs(del_el, next(ids), author, date)
    old_runs[0]._r.addprevious(del_el)
    for r in old_runs:
        for tt in r._r.findall(qn("w:t")):
            tt.tag = qn("w:delText")
        del_el.append(r._r)


@tool
def word_suggest_edits(filename: str, edits: list[dict], author: str = "AI") -> str:
    """Word文書に「変更履歴(見え消し)」として修正を提案する。直接上書きせず、元の文には
    取り消し線、提案文には下線が付き、ユーザーがWordで承認/却下できる。校正・添削・推敲の
    提案に使う(勝手に確定させたくない修正はこちら、確定してよい修正は word_edit_paragraph)。
    editsの各要素は {"index": 段落番号, "new_text": "提案する新しい本文"}(本文を提案文に置換) か
    {"index": 段落番号, "delete": true}(その段落の削除を提案)。indexはword_readの段落番号。
    1つの段落に二重に提案はできない(既に提案済みの段落はスキップする)。編集前にword_readで現状を読むこと。"""
    doc, path = _open_docx(filename)
    paras = list(doc.paragraphs)
    ids = _start_revision_id(doc)
    date = _now_iso()
    ok = 0
    failures: list[str] = []
    seen: set[int] = set()
    for e in edits:
        idx = e.get("index")
        if not isinstance(idx, int) or idx < 0 or idx >= len(paras):
            failures.append(f"[{idx}] 段落番号が範囲外です (0〜{len(paras) - 1})")
            continue
        if idx in seen:
            failures.append(f"[{idx}] 同じ段落が2回指定されています")
            continue
        p = paras[idx]
        if _has_revision(p):
            failures.append(f"[{idx}] すでに変更提案が入っている段落です(スキップ)")
            continue
        if e.get("delete"):
            _tracked_delete(p, ids, author, date)
        elif e.get("new_text") is not None:
            _tracked_replace(p, str(e["new_text"]), ids, author, date)
        else:
            failures.append(f"[{idx}] new_text か delete を指定してください")
            continue
        seen.add(idx)
        ok += 1
    if ok:
        atomic_save(doc.save, path)
    if not ok:
        return "エラー: 1件も提案できませんでした\n" + "\n".join(failures)
    result = f"{filename} に{ok}件の修正を変更履歴として提案しました(Wordで承認/却下できます)"
    if failures:
        result += "\n提案できなかったもの:\n" + "\n".join(failures)
    return result


@tool
def word_add_comments(filename: str, comments: list[dict], author: str = "AI") -> str:
    """Word文書の指定した段落に、レビューコメント(吹き出し)を付ける。本文は変えずに指摘や
    確認事項だけを残したいときに使う(「ここは要確認」等)。commentsの各要素は
    {"index": 段落番号, "text": "コメント本文"}。indexはword_readの段落番号。
    本文のある段落にだけ付けられる(空行には付けられない)。編集前にword_readで現状を読むこと。"""
    doc, path = _open_docx(filename)
    paras = list(doc.paragraphs)
    initials = (author[:2] or "AI").upper()
    ok = 0
    failures: list[str] = []
    for c in comments:
        idx = c.get("index")
        text = c.get("text")
        if not isinstance(idx, int) or idx < 0 or idx >= len(paras):
            failures.append(f"[{idx}] 段落番号が範囲外です (0〜{len(paras) - 1})")
            continue
        if not text or not str(text).strip():
            failures.append(f"[{idx}] コメント本文(text)が空です")
            continue
        runs = paras[idx].runs
        if not runs:
            failures.append(f"[{idx}] 本文が無い段落にはコメントを付けられません")
            continue
        doc.add_comment(runs, text=str(text), author=author, initials=initials)
        ok += 1
    if ok:
        atomic_save(doc.save, path)
    if not ok:
        return "エラー: 1件もコメントを付けられませんでした\n" + "\n".join(failures)
    result = f"{filename} に{ok}件のコメントを付けました"
    if failures:
        result += "\n付けられなかったもの:\n" + "\n".join(failures)
    return result


REVIEW_TOOLS = [doc_diff, word_suggest_edits, word_add_comments]
