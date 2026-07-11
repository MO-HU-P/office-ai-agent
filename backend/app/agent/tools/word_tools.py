"""Word (.docx) 操作ツール群"""
import json
from typing import Any, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.run import Run
from langchain_core.tools import tool

from ...atomic import atomic_save
from ...config import resolve_workspace_path
from .inline_format import describe_format, find_keyword_spans, split_runs_at_spans, validate_format_args

STYLE_MAP = {
    "normal": None,
    "title": "Title",
    "h1": "Heading 1",
    "h2": "Heading 2",
    "h3": "Heading 3",
    "h4": "Heading 4",
    "bullet": "List Bullet",
    "number": "List Number",
    "quote": "Quote",
}


def _open(filename: str) -> tuple[Document, str]:
    path = resolve_workspace_path(filename, must_exist=True)
    if path.suffix != ".docx":
        raise ValueError("Word文書は .docx ファイルを指定してください")
    return Document(str(path)), str(path)


@tool
def word_create(filename: str, title: str = "") -> str:
    """新しいWord文書(.docx)を作成する。filenameは必ず.docxで終わること。titleを指定すると表題段落を追加する。"""
    path = resolve_workspace_path(filename)
    if path.suffix != ".docx":
        return "エラー: filenameは .docx で終わる必要があります"
    doc = Document()
    if title:
        doc.add_heading(title, level=0)
    atomic_save(doc.save, path)
    return f"{filename} を作成しました"


@tool
def word_read(filename: str, mode: str = "full") -> str:
    """Word文書の内容を読む。段落番号・スタイル付きで全段落と表を返す。編集前に必ず呼ぶこと。
    mode="outline"にすると表題・見出しの段落だけを返す。長い文書はまずoutlineで構造を把握し、
    必要な箇所だけをfullで読むとよい。"""
    doc, _ = _open(filename)
    if mode == "outline":
        lines = []
        for i, p in enumerate(doc.paragraphs):
            style = p.style.name if p.style else "Normal"
            if style == "Title" or style.startswith("Heading"):
                lines.append(f"[{i}] ({style}) {p.text.strip()}")
        lines.append(f"(全{len(doc.paragraphs)}段落, 表{len(doc.tables)}個)")
        return "\n".join(lines)
    lines = []
    for i, p in enumerate(doc.paragraphs):
        style = p.style.name if p.style else "Normal"
        text = p.text.strip()
        lines.append(f"[{i}] ({style}) {text}" if text else f"[{i}] ({style}) <空行>")
    for ti, table in enumerate(doc.tables):
        lines.append(f"--- 表{ti} ({len(table.rows)}行 x {len(table.columns)}列) ---")
        for row in table.rows[:20]:
            lines.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(lines) if lines else "(空の文書)"


@tool
def word_append(filename: str, text: str, style: str = "normal") -> str:
    """Word文書の末尾に段落を追加する。textに改行(\\n)を含めると複数段落として追加される。
    style: normal / title / h1 / h2 / h3 / h4 / bullet(箇条書き) / number(番号付き) / quote(引用)"""
    doc, path = _open(filename)
    if not text.strip():
        return "エラー: textが空です。追加する本文を指定してください"
    style_name = STYLE_MAP.get(style, None)
    count = 0
    style_failed = False
    for line in text.split("\n"):
        if style == "normal" and not line.strip():
            doc.add_paragraph("")
            continue
        p = doc.add_paragraph(line)
        if style_name:
            try:
                p.style = style_name
            except KeyError:
                style_failed = True
        count += 1
    atomic_save(doc.save, path)
    result = f"{filename} に{count}段落を追加しました (style={style})"
    if style_failed:
        result += f"\n注意: スタイル「{style_name}」がこの文書に無いため、標準スタイルのままです"
    return result


@tool
def word_find(filename: str, query: str = "", style: str = "") -> str:
    """Word文書から条件に合う段落を探し、段落番号を返す。長い文書で修正箇所を探すときに
    全文をword_readで読む代わりに使う。見つかった番号はそのままword_batch_editに渡せる。
    query: 探す文字列(部分一致)。style: スタイル指定(normal/h1/h2/h3/h4/bullet/number/quote または実際のスタイル名)。
    両方指定すると両方に合う段落だけを返す。少なくとも一方は指定すること。"""
    doc, _ = _open(filename)
    if not query and not style:
        return "エラー: query(探す文字列)かstyle(スタイル)の少なくとも一方を指定してください"
    style_name = STYLE_MAP.get(style, style) if style else None
    hits = []
    for i, p in enumerate(doc.paragraphs):
        p_style = p.style.name if p.style else "Normal"
        if style and p_style != style_name:
            continue
        if query and query not in p.text:
            continue
        text = p.text.strip()
        hits.append(f"[{i}] ({p_style}) {text[:60]}" if text else f"[{i}] ({p_style}) <空行>")
    if not hits:
        return "条件に合う段落は見つかりませんでした"
    lines = hits[:50]
    if len(hits) > 50:
        lines.append(f"...ほか{len(hits) - 50}件")
    lines.append(f"(計{len(hits)}段落が該当)")
    return "\n".join(lines)


def _apply_para_style(p, style: str) -> str | None:
    """段落にスタイル(normal/h1等)を適用する。失敗時はエラーメッセージを返し、成功時はNone。"""
    if style not in STYLE_MAP:
        return f"styleは {' / '.join(STYLE_MAP)} のいずれかを指定してください"
    try:
        p.style = STYLE_MAP[style] or "Normal"
    except KeyError:
        return f"スタイル「{STYLE_MAP[style]}」がこの文書にありません"
    return None


@tool
def word_edit_paragraph(filename: str, index: int, new_text: str = "", delete: bool = False, style: str = "") -> str:
    """既存段落を編集する。indexはword_readで表示される段落番号。delete=Trueでその段落を削除。
    styleを指定すると段落スタイルを変更できる(normal/title/h1/h2/h3/h4/bullet/number/quote)。
    styleだけ指定すれば本文はそのままスタイルだけ変わる(「この段落を見出しにして」に使える)。"""
    doc, path = _open(filename)
    if index < 0 or index >= len(doc.paragraphs):
        return f"エラー: 段落番号 {index} は範囲外です (0〜{len(doc.paragraphs) - 1})"
    p = doc.paragraphs[index]
    if delete:
        p._element.getparent().remove(p._element)
        atomic_save(doc.save, path)
        return f"段落 [{index}] を削除しました"
    if not new_text and not style:
        return "エラー: new_text(新しい本文)・style(スタイル)・delete のいずれかを指定してください"
    done = []
    if new_text:
        # 既存の書式(スタイル)を保ちつつテキストだけ差し替える
        for run in list(p.runs):
            run._element.getparent().remove(run._element)
        p.add_run(new_text)
        done.append("本文")
    if style:
        error = _apply_para_style(p, style)
        if error:
            if done:
                atomic_save(doc.save, path)
            return f"エラー: {error}" if not done else f"段落 [{index}] の本文は更新しましたが、スタイルは失敗: {error}"
        done.append(f"スタイル({style})")
    atomic_save(doc.save, path)
    return f"段落 [{index}] の{'と'.join(done)}を更新しました"


@tool
def word_batch_edit(filename: str, edits: list[dict]) -> str:
    """複数の段落をまとめて編集・削除する。2箇所以上を直すときはword_edit_paragraphを繰り返さず必ずこちらを使う。
    editsの各要素は {"index": 段落番号, "new_text": "新しい本文"} / {"index": 段落番号, "delete": true} /
    {"index": 段落番号, "style": "h1"}(スタイル変更: normal/title/h1/h2/h3/h4/bullet/number/quote。new_textと併用可)。
    indexはword_readで表示される番号(編集前の番号のままでよい)。一部が失敗しても残りは適用される。"""
    doc, path = _open(filename)
    paras = list(doc.paragraphs)  # 削除で番号がずれないよう、編集前の番号で対象を確定しておく
    done: set[int] = set()
    ok_count = 0
    failures: list[str] = []
    for e in edits:
        idx = e.get("index")
        if not isinstance(idx, int) or idx < 0 or idx >= len(paras):
            failures.append(f"[{idx}] 段落番号が範囲外です (0〜{len(paras) - 1})")
            continue
        if idx in done:
            failures.append(f"[{idx}] 同じ段落が2回指定されています")
            continue
        p = paras[idx]
        if e.get("delete"):
            p._element.getparent().remove(p._element)
        elif e.get("new_text") is not None or e.get("style"):
            if e.get("new_text") is not None:
                for run in list(p.runs):
                    run._element.getparent().remove(run._element)
                p.add_run(str(e["new_text"]))
            if e.get("style"):
                error = _apply_para_style(p, str(e["style"]))
                if error:
                    failures.append(f"[{idx}] {error}")
                    if e.get("new_text") is None:
                        continue  # スタイルのみの指定で失敗したら、この段落は未適用扱い
        else:
            failures.append(f"[{idx}] new_text / style / delete のいずれかを指定してください")
            continue
        done.add(idx)
        ok_count += 1
    if ok_count:
        atomic_save(doc.save, path)
    if not ok_count:
        return "エラー: 1件も適用できませんでした\n" + "\n".join(failures)
    result = f"{filename} の{ok_count}段落を更新しました"
    if failures:
        result += "\n適用できなかったもの:\n" + "\n".join(failures)
    return result


_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}
_ALIGN_NAMES = {v: k for k, v in _ALIGN_MAP.items()}


def _dump_one_style(style) -> dict:
    """スタイル1つ分の書式を辞書にする。未設定の属性は継承元(base_style)をたどって解決する。"""
    d: dict = {}
    seen = []
    s = style
    while s is not None and s not in seen and len(seen) < 8:
        seen.append(s)
        f = s.font
        if "font" not in d:
            # 日本語フォント(eastAsia)を優先し、なければ欧文フォント名を使う
            rpr = s.element.rPr
            ea = rpr.rFonts.get(qn("w:eastAsia")) if rpr is not None and rpr.rFonts is not None else None
            if ea or f.name:
                d["font"] = ea or f.name
        if "size" not in d and f.size is not None:
            d["size"] = f.size.pt
        if "bold" not in d and f.bold is not None:
            d["bold"] = f.bold
        if "italic" not in d and f.italic is not None:
            d["italic"] = f.italic
        if "color" not in d and f.color is not None and f.color.rgb is not None:
            d["color"] = f"#{f.color.rgb}"
        pf = getattr(s, "paragraph_format", None)
        if pf is not None:
            if "align" not in d and pf.alignment is not None and pf.alignment in _ALIGN_NAMES:
                d["align"] = _ALIGN_NAMES[pf.alignment]
            if "space_before" not in d and pf.space_before is not None:
                d["space_before"] = pf.space_before.pt
            if "space_after" not in d and pf.space_after is not None:
                d["space_after"] = pf.space_after.pt
        s = s.base_style
    return d


@tool
def word_dump_style(filename: str) -> str:
    """お手本のWord文書から、見出し構造(outline)とスタイル設定(styles: フォント・サイズ・太字・色・
    配置・段落前後の間隔)をJSONで取り出す。「この文書と同じ体裁で作って」と言われたら、
    まずこれでお手本を調べ、styles部分をword_apply_styleで新しい文書に適用してから本文を書くこと。"""
    doc, _ = _open(filename)
    used: list[str] = []
    outline: list[str] = []
    for p in doc.paragraphs:
        name = p.style.name if p.style else "Normal"
        if name not in used:
            used.append(name)
        if name == "Title" or name.startswith("Heading"):
            outline.append(f"({name}) {p.text.strip()}")
    styles = {}
    for name in used:
        try:
            styles[name] = _dump_one_style(doc.styles[name])
        except KeyError:
            continue
    if len(outline) > 50:
        outline = outline[:50] + [f"...ほか{len(outline) - 50}件"]
    return json.dumps({"styles": styles, "outline": outline}, ensure_ascii=False, indent=1)


@tool
def word_apply_style(filename: str, styles: dict) -> str:
    """word_dump_styleで取り出したスタイル設定をWord文書に適用し、お手本と同じ体裁にする。
    stylesはスタイル名→設定の辞書。設定できる項目: font(フォント名) / size(pt) / bold / italic /
    color("#RRGGBB") / align(left・center・right・justify) / space_before・space_after(段落前後の間隔pt)。
    例: {"Heading 1": {"font": "游ゴシック", "size": 14, "bold": true, "color": "#1A73E8"}}
    適用後は既存段落にも、word_appendで追加する段落にも同じ体裁が効く。"""
    doc, path = _open(filename)
    ok_count = 0
    failures: list[str] = []
    for name, conf in styles.items():
        if not isinstance(conf, dict):
            failures.append(f"「{name}」の設定が辞書形式ではありません")
            continue
        try:
            style = doc.styles[name]
        except KeyError:
            failures.append(f"「{name}」というスタイルはこの文書にありません")
            continue
        try:
            f = style.font
            if conf.get("font"):
                font_name = str(conf["font"])
                f.name = font_name
                # 日本語文字にはeastAsiaフォントが使われるため、両方に同じフォントを設定する
                rpr = style.element.get_or_add_rPr()
                rfonts = rpr.rFonts if rpr.rFonts is not None else rpr.get_or_add_rFonts()
                rfonts.set(qn("w:eastAsia"), font_name)
            if conf.get("size") is not None:
                f.size = Pt(float(conf["size"]))
            if conf.get("bold") is not None:
                f.bold = bool(conf["bold"])
            if conf.get("italic") is not None:
                f.italic = bool(conf["italic"])
            if conf.get("color"):
                f.color.rgb = RGBColor.from_string(str(conf["color"]).lstrip("#").upper())
            pf = style.paragraph_format
            if conf.get("align") in _ALIGN_MAP:
                pf.alignment = _ALIGN_MAP[conf["align"]]
            if conf.get("space_before") is not None:
                pf.space_before = Pt(float(conf["space_before"]))
            if conf.get("space_after") is not None:
                pf.space_after = Pt(float(conf["space_after"]))
            ok_count += 1
        except (ValueError, TypeError) as e:
            failures.append(f"「{name}」の適用に失敗しました: {e}")
    if ok_count:
        atomic_save(doc.save, path)
    if not ok_count:
        return "エラー: 1つもスタイルを適用できませんでした\n" + "\n".join(failures)
    result = f"{filename} に{ok_count}個のスタイルを適用しました"
    if failures:
        result += "\n適用できなかったもの:\n" + "\n".join(failures)
    return result


# Wordの蛍光ペンは固定の16色しか使えないため、"#RRGGBB"指定を最も近い色に丸めて適用する
_HIGHLIGHT_PALETTE = {
    WD_COLOR_INDEX.YELLOW: (0xFF, 0xFF, 0x00),
    WD_COLOR_INDEX.BRIGHT_GREEN: (0x00, 0xFF, 0x00),
    WD_COLOR_INDEX.TURQUOISE: (0x00, 0xFF, 0xFF),
    WD_COLOR_INDEX.PINK: (0xFF, 0x00, 0xFF),
    WD_COLOR_INDEX.RED: (0xFF, 0x00, 0x00),
    WD_COLOR_INDEX.BLUE: (0x00, 0x00, 0xFF),
    WD_COLOR_INDEX.GREEN: (0x00, 0x80, 0x00),
    WD_COLOR_INDEX.DARK_YELLOW: (0x80, 0x80, 0x00),
    WD_COLOR_INDEX.TEAL: (0x00, 0x80, 0x80),
    WD_COLOR_INDEX.DARK_BLUE: (0x00, 0x00, 0x80),
    WD_COLOR_INDEX.VIOLET: (0x80, 0x00, 0x80),
    WD_COLOR_INDEX.DARK_RED: (0x80, 0x00, 0x00),
    WD_COLOR_INDEX.GRAY_25: (0xC0, 0xC0, 0xC0),
    WD_COLOR_INDEX.GRAY_50: (0x80, 0x80, 0x80),
    WD_COLOR_INDEX.BLACK: (0x00, 0x00, 0x00),
    WD_COLOR_INDEX.WHITE: (0xFF, 0xFF, 0xFF),
}


def _nearest_highlight(hexv: str):
    """"RRGGBB"を蛍光ペンパレットの最寄り色に丸め、(色インデックス, その色の"#RRGGBB")を返す。"""
    r, g, b = (int(hexv[i:i + 2], 16) for i in (0, 2, 4))
    index, (pr, pg, pb) = min(
        _HIGHLIGHT_PALETTE.items(),
        key=lambda kv: (kv[1][0] - r) ** 2 + (kv[1][1] - g) ** 2 + (kv[1][2] - b) ** 2,
    )
    return index, f"#{pr:02X}{pg:02X}{pb:02X}"


def _iter_all_paragraphs(doc):
    """本文と表のセル内の全段落を順に返す(word_readと同じく1段目の表まで)。"""
    yield from doc.paragraphs
    seen: set[int] = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if id(cell._tc) in seen:  # 結合セルは同じ実体が繰り返し現れるため1回だけ
                    continue
                seen.add(id(cell._tc))
                yield from cell.paragraphs


@tool
def word_format_text(
    filename: str,
    keywords: list[str],
    color: str = "",
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    underline: Optional[bool] = None,
    font_size: Optional[float] = None,
    font: str = "",
    highlight: str = "",
) -> str:
    """Word文書内の特定の語句(キーワード)だけに文字書式を付ける。「重要な用語を赤字にして」
    「◯◯に蛍光ペンを引いて」のような、段落の一部の文字だけの色付け・強調・サイズやフォントの変更に使う
    (段落全体のスタイル変更はword_edit_paragraph、文書全体の体裁はword_apply_styleを使う)。
    文書全体(表の中も含む)からkeywordsの全出現箇所を探して同じ書式を適用する。
    colorは"#RRGGBB"形式(赤字なら"#FF0000")。bold/italic/underlineはtrueで付け、falseで外す。
    font_sizeは文字サイズ(pt)、fontはフォント名(例: "游ゴシック")。
    highlightは蛍光ペン風の背景色で"#RRGGBB"形式(黄色なら"#FFFF00")。Wordの蛍光ペンは
    固定の16色のみのため、指定に最も近い色に丸められる。指定した項目だけが変更される。"""
    keywords, error = validate_format_args(keywords, color, bold, italic, underline, font_size, font, highlight)
    if error:
        return error
    doc, path = _open(filename)
    rgb = RGBColor.from_string(color.lstrip("#").upper()) if color else None
    hl_index = hl_hex = None
    if highlight:
        hl_index, hl_hex = _nearest_highlight(highlight.lstrip("#").upper())
    hit_count = 0
    para_count = 0
    for p in _iter_all_paragraphs(doc):
        runs = list(p.runs)
        spans = find_keyword_spans("".join(r.text for r in runs), keywords)
        if not spans:
            continue
        for el in split_runs_at_spans([r._element for r in runs], spans):
            run = Run(el, p)  # キーワード部分のrunにだけ書式を適用する
            if rgb is not None:
                run.font.color.rgb = rgb
            if bold is not None:
                run.bold = bold
            if italic is not None:
                run.italic = italic
            if underline is not None:
                run.underline = underline
            if font_size is not None:
                run.font.size = Pt(float(font_size))
            if font:
                run.font.name = font
                # 日本語文字にはeastAsiaフォントが使われるため、両方に同じフォントを設定する
                rpr = el.get_or_add_rPr()
                rfonts = rpr.rFonts if rpr.rFonts is not None else rpr.get_or_add_rFonts()
                rfonts.set(qn("w:eastAsia"), font)
            if hl_index is not None:
                run.font.highlight_color = hl_index
        hit_count += len(spans)
        para_count += 1
    if not hit_count:
        return f"「{'」「'.join(keywords)}」は {filename} に見つかりませんでした。word_readで実際の表記を確認してください"
    atomic_save(doc.save, path)
    result = (f"{filename} の{para_count}段落・{hit_count}箇所に書式"
              f"({describe_format(color, bold, italic, underline, font_size, font, hl_hex or '')})を適用しました")
    if hl_hex and hl_hex.lstrip("#") != highlight.lstrip("#").upper():
        result += f"(蛍光ペンはWordで使える近い色 {hl_hex} に丸めています)"
    return result


@tool
def word_add_table(filename: str, rows: list[list[Any]], header: bool = True) -> str:
    """Word文書の末尾に表を追加する。rowsは2次元配列(行のリスト)で、セルは文字列・数値どちらでもよい。
    header=Trueなら1行目を太字ヘッダーにする。"""
    doc, path = _open(filename)
    if not rows or not rows[0]:
        return "エラー: rowsが空です"
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    for ri, row in enumerate(rows):
        for ci in range(n_cols):
            val = str(row[ci]) if ci < len(row) and row[ci] is not None else ""
            cell = table.cell(ri, ci)
            cell.text = val
            if header and ri == 0:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.size = Pt(10.5)
    atomic_save(doc.save, path)
    return f"{filename} に {len(rows)}x{n_cols} の表を追加しました"


WORD_TOOLS = [
    word_create,
    word_read,
    word_find,
    word_append,
    word_edit_paragraph,
    word_batch_edit,
    word_format_text,
    word_add_table,
    word_dump_style,
    word_apply_style,
]
