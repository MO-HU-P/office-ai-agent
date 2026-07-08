"""変更差分ビュー(services/changes)のテスト。Wordは段落単位、Excelはセル単位。"""
from docx import Document
from openpyxl import Workbook, load_workbook

from app.atomic import atomic_save
from app.services import changes, history


def _diff_texts(result, op):
    return [line["text"] for line in result["lines"] if line["op"] == op]


def test_docx_paragraph_diff(ws):
    path = ws / "doc.docx"
    doc = Document()
    doc.add_paragraph("こんにちは")
    doc.add_paragraph("さようなら")
    atomic_save(doc.save, path)

    history.begin_turn()
    doc2 = Document(str(path))
    doc2.paragraphs[1].runs[0].text = "また明日"
    atomic_save(doc2.save, path)
    history.end_turn()

    result = changes.build_changes("doc.docx")
    assert result["available"] is True
    assert "さようなら" in _diff_texts(result, "del")
    assert "また明日" in _diff_texts(result, "add")
    assert result["added"] == 1 and result["removed"] == 1


def test_xlsx_cell_level_diff(ws):
    path = ws / "book.xlsx"
    wb = Workbook()
    sh = wb.active
    sh["A1"] = "商品"
    sh["B1"] = 100
    atomic_save(wb.save, path)

    history.begin_turn()
    wb2 = load_workbook(str(path))
    wb2.active["B1"] = 200
    wb2.active["C1"] = "新セル"
    atomic_save(wb2.save, path)
    history.end_turn()

    result = changes.build_changes("book.xlsx")
    assert result["available"] is True
    dels = _diff_texts(result, "del")
    adds = _diff_texts(result, "add")
    # セル番地つきで、変わったセルだけが差分になる(A1は差分に出ない)
    assert any("B1: 100" in t for t in dels)
    assert any("B1: 200" in t for t in adds)
    assert any("C1: 新セル" in t for t in adds)
    assert not any("A1" in t for t in dels + adds)


def test_new_file_has_no_changes_view(ws):
    doc = Document()
    doc.add_paragraph("新規")
    atomic_save(doc.save, ws / "new.docx")
    result = changes.build_changes("new.docx")
    assert result["available"] is False
    assert "バックアップ" in result["reason"]


def test_identical_content_reports_no_changes(ws):
    path = ws / "doc.docx"
    doc = Document()
    doc.add_paragraph("同じ")
    atomic_save(doc.save, path)
    # 内容は同じまま保存し直す(書式だけの変更などを想定)
    doc2 = Document(str(path))
    atomic_save(doc2.save, path)
    result = changes.build_changes("doc.docx")
    assert result["available"] is False
    assert "変更はありません" in result["reason"]


def test_unsupported_extension(ws):
    (ws / "img.png").write_bytes(b"\x89PNG")
    result = changes.build_changes("img.png")
    assert result["available"] is False
